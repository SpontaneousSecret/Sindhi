from docx import Document
import os

def update_contract():
    docx_path = r'docs/Contract.docx'
    
    if not os.path.exists(docx_path):
        print(f"Error: {docx_path} not found.")
        return

    doc = Document(docx_path)

    # 1. Update Section 2.2 Checkout clause
    found_checkout = False
    for para in doc.paragraphs:
        if "restricted to Cash on Delivery (COD) only" in para.text:
            para.text = para.text.replace(
                "restricted to Cash on Delivery (COD) only. No online payment gateway integration is included in this phase.",
                "includes Cash on Delivery (COD) and Online Payments via Axis Bank Payment Gateway integration."
            )
            found_checkout = True
            print("Updated Checkout clause.")
            break
    
    if not found_checkout:
        print("Warning: Could not find strict match for checkout clause. Attempting fuzzy match...")
        for para in doc.paragraphs:
            if "Cash on Delivery (COD)" in para.text and "only" in para.text:
                para.text = "Checkout process includes Cash on Delivery (COD) and Online Payments via Axis Bank Payment Gateway integration."
                found_checkout = True
                print("Updated Checkout clause (fuzzy).")
                break

    # 2. Append Section 6.4 Payment Gateway Terms and Liability
    # Find the end of Section 6 or just append before the signature section
    signature_found = False
    insert_index = len(doc.paragraphs)
    for i, para in enumerate(doc.paragraphs):
        if "IN WITNESS WHEREOF" in para.text or "Service Provider" in para.text:
            insert_index = i
            signature_found = True
            break
    
    new_clauses = [
        "",
        "4. Payment Gateway Terms and Liability:",
        "    - Axis Bank Integration: The Service Provider shall integrate the Axis Bank Payment Gateway for the processing of online transactions.",
        "    - Gateway Removal and Modification: In the event that the Client requests the removal or decoupling of the payment gateway from the platform at any future date, a specialized engineering and reconfiguration fee of INR 3,500 to INR 4,000 shall be payable by the Client. Such removal process necessitates a scheduled platform suspension of 7 to 10 business days to ensure structural integrity and security compliance.",
        "    - Limitation of Liability: The Service Provider’s responsibility is limited to the technical integration of the payment gateway. The Service Provider shall not be held liable for any disputes, delays, processing errors, settlement failures, or transaction issues arising within the banking infrastructure. Any and all matters related to payment clearances and fund settlements are the exclusive responsibility of Axis Bank.",
        ""
    ]

    # Insert before signature or append
    for text in reversed(new_clauses):
        new_p = doc.paragraphs[insert_index-1].insert_paragraph_before(text)
        # Optional: Add formatting if needed, but plain text is safer for matching existing style
        
    print("Added Payment Gateway Terms and Liability clauses.")

    # Save the document
    doc.save(docx_path)
    print(f"Successfully updated {docx_path}")

if __name__ == "__main__":
    update_contract()
